from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "DashCart_Database_Systems_Report.docx"
ARCHITECTURE_IMAGE = ROOT / "architecture.jpeg"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 102, 122)
LIGHT_FILL = "E8EEF5"
LIGHT_ALT = "F8FAFD"


SCHEMA_ROWS = [
    ("users_user", "Application customer/staff identity used by the business domain", "PK user_id"),
    ("users_accountprofile", "Maps Django auth users to roles and assigned store", "FK user -> auth_user, FK store -> stores_store"),
    ("stores_store", "Business store entity", "PK store_id"),
    ("stores_warehouse", "Warehouse for each store", "PK warehouse_id, FK store -> stores_store"),
    ("products_brand", "Product brand master", "PK brand_id"),
    ("products_category", "Category master", "PK category_id"),
    ("products_product", "Product catalog", "PK product_id, FK brand -> products_brand"),
    ("products_productcategory", "Many-to-many bridge between products and categories", "FK product, FK category"),
    ("products_productfavorite", "Wishlist/favorite bridge between customer and product", "FK customer, FK product"),
    ("inventory_inventory", "Stock quantity of a product in a warehouse", "FK product, FK warehouse"),
    ("stock_movements", "Audit trail of inventory changes", "PK movement_id, FK product, FK warehouse"),
    ("suppliers_supplier", "Supplier master data", "PK supplier_id"),
    ("supplier_products", "Many-to-many bridge between suppliers and products", "FK supplier, FK product"),
    ("suppliers_supplierrequest", "Procurement request raised by staff", "PK request_id, FK supplier, FK requested_by"),
    ("discounts_discount", "Discount definition optionally scoped to a store", "PK discount_id, FK store"),
    ("discounts_productdiscount", "Bridge assigning discounts to products", "FK product, FK discount"),
    ("cart_cart", "Active shopping cart for a user", "FK user"),
    ("cart_cartitem", "Items in the cart with chosen warehouse", "FK cart, FK product, FK warehouse"),
    ("orders_order", "Customer order header", "PK order_id, FK user"),
    ("orders_orderitem", "Order line items", "FK order, FK product"),
    ("orders_orderitemallocation", "Warehouse allocation for each order item", "FK order_item, FK warehouse"),
    ("payments_payment", "Payment record per order", "PK payment_id, O2O order"),
    ("payments_refund", "Refund record tied to a payment", "O2O payment"),
    ("reviews", "Customer review for a product", "PK review_id, FK product, FK customer"),
]

RELATION_ROWS = [
    ("Store", "Warehouse", "1:M", "A store may contain multiple warehouses."),
    ("Brand", "Product", "1:M", "Each product belongs to one brand."),
    ("Product", "Category", "M:N", "Implemented through products_productcategory."),
    ("Product", "Inventory", "1:M", "One product can exist in many warehouses."),
    ("Warehouse", "Inventory", "1:M", "One warehouse can hold many products."),
    ("Supplier", "Product", "M:N", "Implemented through supplier_products."),
    ("User", "Cart", "1:1", "A customer owns one active cart."),
    ("Cart", "CartItem", "1:M", "A cart can contain many items."),
    ("User", "Order", "1:M", "A customer can place multiple orders."),
    ("Order", "OrderItem", "1:M", "Each order contains multiple order items."),
    ("OrderItem", "Warehouse", "M:N", "Implemented through orders_orderitemallocation."),
    ("Order", "Payment", "1:1", "Each order has one payment record."),
    ("Payment", "Refund", "1:0..1", "A payment may have one refund record."),
    ("Product", "Review", "1:M", "A product can receive many reviews."),
    ("User", "Review", "1:M", "A customer can review multiple products."),
    ("Store", "Discount", "1:M", "Discounts may be scoped to a store."),
    ("Discount", "Product", "M:N", "Implemented through discounts_productdiscount."),
]

NORMALIZATION_ROWS = [
    ("1NF", "Atomic fields only", "Product, store, supplier, and order data are stored in separate columns without repeating groups."),
    ("2NF", "Full dependency on whole key", "Bridge tables such as products_productcategory and supplier_products use only attributes that depend on the full relationship key."),
    ("3NF", "No transitive dependency", "Store details are stored in stores_store rather than repeated in warehouse, inventory, or order tables."),
    ("3NF", "Role/profile split", "Authentication data stays in Django auth_user while business role/store assignment is kept in users_accountprofile."),
    ("3NF", "Operational separation", "Payments, refunds, reviews, discounts, and stock movements are separated from core order/product tables to avoid update anomalies."),
]

QUERY_ROWS = [
    (
        "Low-stock procedure",
        "CALL sp_low_stock_products(5);",
        "Returns products whose total stock across warehouses is below a threshold.",
    ),
    (
        "Total stock function",
        "SELECT p.product_id, p.name, fn_total_stock(p.product_id) AS total_stock FROM products_product p;",
        "Shows the total inventory of each product using a reusable MySQL function.",
    ),
    (
        "Discounted products",
        "SELECT p.product_id, p.name, d.name AS discount_name, d.discount_percent FROM products_product p JOIN discounts_productdiscount pd ON p.product_id = pd.product_id JOIN discounts_discount d ON pd.discount_id = d.discount_id;",
        "Lists products that currently have discounts attached to them.",
    ),
    (
        "Orders with payment totals",
        "SELECT o.order_id, u.name AS customer_name, pay.amount, pay.status FROM orders_order o JOIN users_user u ON o.user_id = u.user_id JOIN payments_payment pay ON pay.order_id = o.order_id;",
        "Combines order header, customer, and payment information for presentation or reporting.",
    ),
    (
        "Correlated rating query",
        "SELECT p.product_id, p.name, (SELECT AVG(r.rating) FROM reviews r WHERE r.product_id = p.product_id) AS avg_rating FROM products_product p;",
        "A correlated subquery that computes the average rating for each product.",
    ),
    (
        "Nested query for active products",
        "SELECT name FROM products_product WHERE product_id IN (SELECT product_id FROM inventory_inventory WHERE quantity > 0);",
        "Nested query used to find products that are currently available in at least one warehouse.",
    ),
]

MYSQL_OBJECT_ROWS = [
    ("Trigger", "trg_order_status_audit", "Automatically logs order status changes into order_status_audit after each update."),
    ("Function", "fn_total_stock", "Computes total stock of a product by summing inventory rows across warehouses."),
    ("Procedure", "sp_low_stock_products", "Returns all products below a requested stock threshold."),
    ("Trigger", "trg_inventory_stock_movement", "Writes stock movement entries whenever inventory quantities change."),
]

MODULE_ROWS = [
    ("Customer module", "Browse products, choose store/warehouse, add to cart, place orders, track orders, manage wishlist, submit reviews."),
    ("Staff module", "Manage products, inventory, suppliers, supplier requests, discounts, and observe sales/low-stock dashboard metrics."),
    ("Database logic", "MySQL functions, procedures, and triggers automate stock calculation, low-stock reporting, and auditing."),
    ("Presentation layer", "React frontend communicates with Django REST endpoints that map directly to relational entities and workflows."),
]

DEMO_ROWS = [
    ("1", "Login as customer", "Show dashboard, products, deals section, and available products."),
    ("2", "Browse discounted products", "Demonstrate the deals strip and open one product detail page."),
    ("3", "Select store and add to cart", "Show that the selected warehouse/store controls the cart item allocation."),
    ("4", "Checkout and view orders", "Explain how order, order item, payment, and allocation tables are populated."),
    ("5", "Submit a review", "Show review creation and how product detail pages display customer reviews."),
    ("6", "Login as staff", "Open inventory, suppliers, discounts, and dashboard KPIs."),
    ("7", "Discuss MySQL objects", "Present triggers, functions, procedures, and one or two sample SQL queries from the report."),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    tr_pr.append(hdr)


def set_paragraph_border_bottom(paragraph, color="D9E2F3", size=8, space=1):
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = p_pr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        p_pr.append(pbdr)
    bottom = pbdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pbdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def style_run(run, *, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Title", 24, INK, 0, 8),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run("DashCart Database Systems Report")
    style_run(footer_run, size=9, color=MUTED)


def add_cover_page(doc):
    doc.add_paragraph("")
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("DATABASE SYSTEMS COURSE PROJECT REPORT")
    style_run(run, size=12, bold=True, color=BLUE)
    kicker.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("DashCart")
    style_run(title_run, size=28, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Retail Inventory, Order, and Supplier Management System")
    style_run(subtitle_run, size=15, color=DARK_BLUE)
    subtitle.paragraph_format.space_after = Pt(8)

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2_run = subtitle2.add_run(
        "Prepared for the Database Systems assignment presentation"
    )
    style_run(subtitle2_run, size=11, color=MUTED, italic=True)
    subtitle2.paragraph_format.space_after = Pt(28)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    widths = [Inches(2.0), Inches(4.0)]
    for row in meta.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    entries = [
        ("Project Type", "Website-backed relational database application"),
        ("Academic Focus", "ERD, relational mapping, normalization, queries, procedures, functions, and triggers"),
        ("Technology Stack", "MySQL (target), Django REST backend, React frontend"),
        ("Team Details", "Add student names, IDs, and section before final submission"),
    ]
    for i, (label, value) in enumerate(entries):
        left, right = meta.rows[i].cells
        left.text = label
        right.text = value
        set_cell_shading(left, LIGHT_FILL)
        for para in left.paragraphs + right.paragraphs:
            for run in para.runs:
                style_run(run, size=10.5, bold=(para is left.paragraphs[0]))

    doc.add_paragraph("")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run(
        "This report is designed to support the live presentation and viva for the Database Systems course."
    )
    style_run(note_run, size=10.5, color=MUTED)

    doc.add_section(WD_SECTION.NEW_PAGE)


def add_section_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)
    set_paragraph_border_bottom(p)


def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"• {item}")
        style_run(run, size=11)


def add_table(doc, headers, rows, col_widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        hdr_cells[idx].width = col_widths[idx]
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(hdr_cells[idx], LIGHT_FILL)
        for run in hdr_cells[idx].paragraphs[0].runs:
            style_run(run, size=10.5, bold=True, color=INK)
    set_repeat_table_header(table.rows[0])

    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = str(value)
            cells[idx].width = col_widths[idx]
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index % 2 == 1:
                set_cell_shading(cells[idx], LIGHT_ALT)
            for run in cells[idx].paragraphs[0].runs:
                style_run(run, size=10)
    return table


def add_architecture_image(doc):
    if not ARCHITECTURE_IMAGE.exists():
        return
    doc.add_paragraph(style="Heading 2").add_run("Repository Architecture Reference")
    p = doc.add_paragraph(
        "The repository already contains an architecture image. It is included below as a visual reference for the current application structure."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(ARCHITECTURE_IMAGE), width=Inches(5.8))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("Figure 1. Existing architecture reference from the project repository.")
    style_run(r, size=10, color=MUTED, italic=True)


def build_doc():
    doc = Document()
    configure_document(doc)
    add_cover_page(doc)

    add_section_heading(doc, "1. Problem Statement")
    doc.add_paragraph(
        "DashCart is a retail management website that allows customers to browse products, select a store, add items to a cart, place orders, manage a wishlist, and submit reviews. Staff users manage products, warehouses, inventory, suppliers, discounts, and order fulfillment. The project demonstrates how a market-oriented retail problem can be transformed into a structured relational database design and then connected to a usable web application."
    )

    add_section_heading(doc, "2. Project Objectives")
    add_bullet_list(
        doc,
        [
            "Design a practical relational database for a multi-store retail business.",
            "Map business entities such as products, warehouses, orders, suppliers, payments, and reviews into normalized relational tables.",
            "Implement key database operations including nested queries, correlated queries, stored procedures, functions, and triggers.",
            "Build a front-end and back-end workflow that shows how database concepts are used in a realistic website.",
            "Prepare the system for a classroom presentation that emphasizes database design decisions rather than only UI behavior.",
        ],
    )

    add_section_heading(doc, "3. Scope and User Roles")
    add_table(
        doc,
        ["Module", "Scope"],
        MODULE_ROWS,
        [Inches(1.7), Inches(4.8)],
    )

    add_section_heading(doc, "4. Entity-Relationship Design Summary")
    doc.add_paragraph(
        "The core ERD of DashCart is centered around product movement across stores and warehouses. Products belong to brands, may appear in multiple categories, may be stocked in multiple warehouses, and may receive discounts and reviews. Orders, carts, and favorites connect users to products, while supplier tables support procurement and replenishment."
    )
    add_table(
        doc,
        ["Entity A", "Entity B", "Cardinality", "Meaning"],
        RELATION_ROWS,
        [Inches(1.25), Inches(1.35), Inches(0.9), Inches(3.0)],
    )
    add_architecture_image(doc)

    add_section_heading(doc, "5. Relational Schema Overview")
    doc.add_paragraph(
        "The following table summarizes the main schema used in the application. Bridge tables are used whenever many-to-many relationships appear in the conceptual model."
    )
    add_table(
        doc,
        ["Table", "Purpose", "Key Relationships"],
        SCHEMA_ROWS,
        [Inches(2.0), Inches(2.85), Inches(1.65)],
    )

    add_section_heading(doc, "6. Normalization Discussion")
    doc.add_paragraph(
        "The database is organized to reduce redundancy and update anomalies. Instead of storing product, store, order, and payment details in a single wide table, the design separates each concern into its own relation and uses foreign keys to preserve integrity."
    )
    add_table(
        doc,
        ["Normal Form", "Design Rule", "How DashCart Satisfies It"],
        NORMALIZATION_ROWS,
        [Inches(0.9), Inches(1.6), Inches(4.0)],
    )

    add_section_heading(doc, "7. MySQL Functions, Procedures, and Triggers")
    doc.add_paragraph(
        "Because the course emphasizes practical MySQL concepts, DashCart includes SQL objects that go beyond basic tables. These objects are defined in sql/mysql_objects.sql and sql/mysql_existing_table_objects.sql and can be applied using the Django management command that loads the SQL files."
    )
    add_table(
        doc,
        ["Object Type", "Name", "Purpose"],
        MYSQL_OBJECT_ROWS,
        [Inches(1.1), Inches(2.0), Inches(3.4)],
    )
    add_paragraph = doc.add_paragraph(
        "The Django backend also contains helper methods in backend/mysql_routines.py that call MySQL procedures and functions such as cart checkout, stock movement creation, refund requests, effective price calculation, and total inventory computation."
    )
    add_paragraph.paragraph_format.space_after = Pt(8)

    add_section_heading(doc, "8. Sample SQL Queries for Presentation")
    doc.add_paragraph(
        "These sample SQL statements can be shown directly in the presentation to demonstrate database logic. They cover functions, procedures, joins, nested queries, and correlated queries."
    )
    add_table(
        doc,
        ["Query Type", "SQL Example", "What It Demonstrates"],
        QUERY_ROWS,
        [Inches(1.2), Inches(3.8), Inches(1.5)],
    )

    add_section_heading(doc, "9. How the Website Uses the Database")
    add_bullet_list(
        doc,
        [
            "The dashboard calculates KPIs such as total sales, low stock counts, cart totals, and product availability using relational joins and aggregates.",
            "The product pages combine products, brands, categories, discounts, reviews, and inventory availability into a customer-friendly catalog.",
            "The cart workflow stores the chosen warehouse for each cart item, showing that stock is store-specific rather than globally available.",
            "The checkout workflow generates an order header, order items, payment records, and warehouse allocations.",
            "The suppliers and inventory modules support procurement and stock replenishment, which reflects real-world retail operations.",
        ],
    )

    add_section_heading(doc, "10. Suggested Presentation Flow")
    add_table(
        doc,
        ["Step", "Live Demo Action", "Talking Point"],
        DEMO_ROWS,
        [Inches(0.45), Inches(2.15), Inches(3.9)],
    )

    add_section_heading(doc, "11. Strengths of the Project")
    add_bullet_list(
        doc,
        [
            "Covers both customer-facing and staff-facing business processes.",
            "Contains enough entities and relationships to justify a substantial ERD.",
            "Uses bridge tables and warehouse-based stock logic, which strengthens the relational design.",
            "Includes MySQL functions, procedures, and triggers for academic depth.",
            "Demonstrates how normalized tables support a practical web application instead of an isolated SQL script.",
        ],
    )

    add_section_heading(doc, "12. Limitations and Future Improvements")
    add_bullet_list(
        doc,
        [
            "A finalized graphical ERD should be added before submission if the team wants a stronger visual database explanation.",
            "The local development environment often uses SQLite for convenience, while the course presentation should emphasize the MySQL design and SQL objects.",
            "Additional analytics queries can be added, such as top-selling products, supplier performance, and monthly revenue trends.",
            "Trigger-driven auditing can be expanded to cover payment status changes and discount activity.",
            "A formal appendix with screenshots from the live website can be added if the instructor values front-end evidence.",
        ],
    )

    add_section_heading(doc, "13. Conclusion")
    doc.add_paragraph(
        "DashCart is a suitable Database Systems course project because it demonstrates the full path from problem definition to ERD thinking, normalized relational design, SQL object creation, and web-based execution. The project is especially strong for presentation when the team clearly explains the schema, the many-to-many bridge tables, the warehouse-based stock model, and the MySQL procedures/functions/triggers that automate business rules."
    )

    doc.add_paragraph("")
    final_note = doc.add_paragraph()
    note_run = final_note.add_run(
        "Final edit before submission: replace the placeholder team metadata on the cover page with actual student names, IDs, and section details."
    )
    style_run(note_run, size=10.5, color=MUTED, italic=True)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
